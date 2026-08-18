# Instalar librería: pip install python-docx
from docx import Document

doc = Document()

# Título
doc.add_heading('Reporte de Análisis y Diseño del Sistema de Gestión de Trámites Ciudadanos', level=1)

# Introducción
doc.add_heading('Introducción', level=2)
doc.add_paragraph(
    "El propósito de esta práctica es analizar una problemática real de la administración pública "
    "y transformarla en un diseño orientado a objetos en Java que reemplace los esquemas tradicionales "
    "basados en hojas de cálculo y papel. La gestión manual genera problemas recurrentes como duplicidad "
    "de registros, falta de seguimiento en tiempo real y lentitud en la atención.\n\n"
    "Para resolver estas necesidades, se toma como caso de estudio inicial el Trámite de Licencia de Conducir. "
    "Este proceso requiere validar la identidad del solicitante mediante CURP e identificación oficial, registrar pagos, "
    "coordinar citas, capturar datos biométricos y evaluar exámenes teóricos y prácticos antes de la dictaminación de un servidor público."
)

# Desarrollo
doc.add_heading('Desarrollo', level=2)
doc.add_heading('Identificación de Clases y Objetos', level=3)
p_clases = doc.add_paragraph()
p_clases.add_run('• Persona: ').bold = True
p_clases.add_run('Clase base para centralizar los datos biográficos generales.\n')
p_clases.add_run('• Ciudadano y ServidorPublico: ').bold = True
p_clases.add_run('Clases especializadas que representan a los actores del sistema.\n')
p_clases.add_run('• SolicitudLicencia: ').bold = True
p_clases.add_run('Núcleo del proceso donde se vinculan el ciudadano, los pagos, documentos y exámenes.\n')
p_clases.add_run('• Soporte: ').bold = True
p_clases.add_run('TipoLicencia, Requisito, DocumentoAdjunto, Pago y ExamenConduccion para validación técnica y financiera.')

# Tabla de Atributos y Métodos
doc.add_heading('Definición de Atributos y Métodos', level=3)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Clase'
hdr_cells[1].text = 'Atributos (Tipos de datos)'
hdr_cells[2].text = 'Métodos principales'

datos = [
    ("Persona", "id (int), nombre (String), apellidos (String), curp (String), correo (String)", "obtenerNombreCompleto(), validarCurp()"),
    ("Ciudadano", "(Hereda de Persona) direccion (String), telefono (String)", "actualizaDatos(), consultarEstado()"),
    ("ServidorPublico", "(Hereda de Persona) numeroEmpleado (String), area (String)", "dictaminarSolicitud(), aprobarDocumento()"),
    ("SolicitudLicencia", "folio (int), estado (String), requisitos (List), examenes (List)", "evaluarElegibilidad(), agregarRequisito()"),
    ("Pago", "idPago (int), monto (double), folioPago (String), estatus (boolean)", "registrarPago(), validarMonto()"),
    ("ExamenConduccion", "idExamen (int), tipo (String), calificacion (double), aprobado (boolean)", "calificar(), esAprobatorio()")
]

for clase, attrs, mets in datos:
    row_cells = table.add_row().cells
    row_cells[0].text = clase
    row_cells[1].text = attrs
    row_cells[2].text = mets

# Conclusiones
doc.add_heading('Conclusiones', level=2)
doc.add_paragraph(
    "• Estructura y Mantenibilidad: La orientación a objetos permite aislar las reglas de negocio de la licencia de conducir en clases cohesivas.\n"
    "• Integración con Base de Datos: La transición de objetos a tablas relacionales mediante JDBC demuestra cómo la abstracción facilita la persistencia.\n"
    "• Trazabilidad: Modelar cada cambio de estado dentro de SolicitudLicencia elimina la opacidad operativa."
)

# Guardar documento
doc.save("Reporte_Tramites_Ciudadanos.docx")
print("Documento guardado exitosamente como 'Reporte_Tramites_Ciudadanos.docx'")
